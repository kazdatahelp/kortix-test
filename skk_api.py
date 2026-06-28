#!/usr/bin/env python3.11
"""SKK Dashboard API v2 — multi-product summary + DOCX report generation.
IBCS-compliant. Reads eoz_full_auth.db."""
import sqlite3, statistics, os, io, datetime
from pathlib import Path
from typing import Optional
from fastapi import FastAPI, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from pydantic import BaseModel

DB_PATH = Path(os.environ.get("EOZ_DB", Path(__file__).parent / "eoz_full_auth.db"))
OUTPUT_DIR = Path(__file__).parent / "output"

app = FastAPI(title="SKK Dashboard API v2", version="2.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

TRU_TYPE_MAP = {1: "товар", 2: "работа", 3: "услуга"}
TRU_TYPE_REV = {"товар": 1, "работа": 2, "услуга": 3}

def get_db() -> sqlite3.Connection:
    db = sqlite3.connect(str(DB_PATH))
    db.row_factory = sqlite3.Row
    db.execute("PRAGMA journal_mode=WAL")
    return db

# ── Helpers ─────────────────────────────────────────────

def _compute_stats(prices: list) -> dict:
    """Compute median, P25, P75, spread, confidence class from price list."""
    n = len(prices)
    if n < 3:
        return {"median": 0, "p25": 0, "p75": 0, "spread": 0, "conf_class": "C (ненадёжно)", "n": n}
    ps = sorted(prices)
    p25 = ps[int(n * 0.25)]
    p50 = statistics.median(ps)
    p75 = ps[int(n * 0.75)]
    spread = ((p75 - p25) / p50 * 100) if p50 > 0 else 0
    if n >= 30 and spread < 150:
        conf = "A (надёжно)"
    elif n >= 8:
        conf = "B (ограниченно)"
    else:
        conf = "C (ненадёжно)"
    return {"median": round(p50, 2), "p25": round(p25, 2), "p75": round(p75, 2),
            "spread": round(spread, 3), "conf_class": conf, "n": n}

# ── Endpoints ───────────────────────────────────────────

@app.get("/api/health")
def health():
    return {"status": "ok", "db": str(DB_PATH), "exists": DB_PATH.exists()}

@app.get("/api/products")
def list_products(
    q: Optional[str] = Query(None),
    tru_type: Optional[int] = Query(None),
    limit: int = Query(200, le=500),
):
    db = get_db()
    try:
        where = ["c.status = 'EXECUTED'"]
        params = []
        if q:
            where.append("(pe.enstru LIKE ? OR pe.name LIKE ?)")
            params.extend([f"%{q}%", f"%{q}%"])
        if tru_type is not None:
            where.append("de.tru_type = ?")
            params.append(tru_type)

        sql = f"""
            SELECT pe.enstru, pe.name, de.tru_type,
                   COUNT(*) as contract_count,
                   COUNT(DISTINCT c.provider_bin) as supplier_count,
                   ROUND(AVG(c.amt / NULLIF(c.qty, 0)), 2) as avg_unit_price,
                   SUM(c.amt) as total_sum,
                   MIN(c.year) as year_from, MAX(c.year) as year_to
            FROM skk_contract c
            JOIN skk_plan_enstru pe ON c.plan_id = pe.plan_id
            LEFT JOIN skk_dim_enstru de ON pe.enstru = de.code
            WHERE {' AND '.join(where)}
            GROUP BY pe.enstru
            ORDER BY contract_count DESC
            LIMIT {limit}
        """
        rows = db.execute(sql, params).fetchall()
        return [{
            "enstru": r["enstru"], "name": r["name"] or "",
            "tru_type": TRU_TYPE_MAP.get(r["tru_type"], "неизвестно"),
            "contract_count": r["contract_count"], "supplier_count": r["supplier_count"],
            "avg_unit_price": r["avg_unit_price"], "total_sum": r["total_sum"],
            "year_from": r["year_from"], "year_to": r["year_to"],
        } for r in rows]
    finally:
        db.close()

@app.get("/api/regions")
def list_regions(enstru: Optional[str] = Query(None)):
    db = get_db()
    try:
        joins, where, params = [], ["c.status = 'EXECUTED'"], []
        if enstru:
            joins.append("JOIN skk_plan_enstru pe ON c.plan_id = pe.plan_id")
            where.append("pe.enstru = ?"); params.append(enstru)
        sql = f"""
            SELECT rk.ab as region_code, rk.full_name_ru as region_name,
                   COUNT(*) as contract_count, COUNT(DISTINCT c.provider_bin) as supplier_count
            FROM skk_contract c
            {' '.join(joins)}
            JOIN skk_ref_kato rk ON CAST(c.delivery_kato_code / 10000000 AS INTEGER) = CAST(rk.ab AS INTEGER)
            WHERE {' AND '.join(where)}
            GROUP BY rk.ab ORDER BY contract_count DESC
        """
        return [dict(r) for r in db.execute(sql, params).fetchall()]
    finally:
        db.close()

@app.get("/api/summary")
def get_summary(
    year_from: int = Query(2024), year_to: int = Query(2026),
    tru_type: Optional[str] = Query(None, description="товар|работа|услуга"),
    exclude_rnu: bool = Query(True),
):
    """Multi-product summary: class distribution, supplier pool, savings estimate."""
    db = get_db()
    try:
        where = ["c.status = 'EXECUTED'", "c.year BETWEEN ? AND ?"]
        params = [year_from, year_to]
        if tru_type:
            where.append("de.tru_type = ?")
            params.append(TRU_TYPE_REV.get(tru_type, 1))

        # Per-ENSTRU stats: count items with valid prices
        sql = f"""
            SELECT pe.enstru, pe.name, de.tru_type,
                   COUNT(*) as n_contracts,
                   COUNT(DISTINCT c.provider_bin) as n_suppliers,
                   SUM(c.amt) as total_sum
            FROM skk_contract c
            JOIN skk_plan_enstru pe ON c.plan_id = pe.plan_id
            LEFT JOIN skk_dim_enstru de ON pe.enstru = de.code
            WHERE {' AND '.join(where)}
              AND c.qty > 0 AND c.amt > 0
            GROUP BY pe.enstru
        """
        positions = db.execute(sql, params).fetchall()

        # Compute class distribution
        class_counts = {"A": 0, "B": 0, "C": 0}
        total_positions = 0
        position_details = []

        for p in positions:
            # Get individual prices for this ENSTRU
            prices = db.execute("""
                SELECT ROUND(c.amt / NULLIF(c.qty, 0), 2) as up
                FROM skk_contract c
                JOIN skk_plan_enstru pe ON c.plan_id = pe.plan_id
                WHERE pe.enstru = ? AND c.status = 'EXECUTED'
                  AND c.year BETWEEN ? AND ? AND c.qty > 0 AND c.amt > 0
                ORDER BY up
            """, [p["enstru"], year_from, year_to]).fetchall()
            price_list = [r["up"] for r in prices if r["up"] and r["up"] > 0]

            if len(price_list) < 3:
                continue

            stats = _compute_stats(price_list)
            total_positions += 1
            cls_key = stats["conf_class"][0]
            class_counts[cls_key] = class_counts.get(cls_key, 0) + 1

            position_details.append({
                "enstru": p["enstru"], "name": p["name"] or "",
                "tru_type": TRU_TYPE_MAP.get(p["tru_type"], "неизвестно"),
                "median_price": stats["median"], "p25": stats["p25"], "p75": stats["p75"],
                "spread_pct": stats["spread"], "conf_class": stats["conf_class"],
                "n_items": stats["n"], "n_suppliers": p["n_suppliers"],
                "total_sum": p["total_sum"],
            })

        # Supplier pool
        supp_sql = f"""
            SELECT COUNT(DISTINCT c.provider_bin) as total,
                   COUNT(DISTINCT CASE WHEN ds.is_rnu = 1 THEN c.provider_bin END) as rnu_count
            FROM skk_contract c
            JOIN skk_plan_enstru pe ON c.plan_id = pe.plan_id
            LEFT JOIN skk_dim_subject ds ON c.provider_bin = ds.bin_iin
            WHERE {' AND '.join(where)}
        """
        supplier_stats = db.execute(supp_sql, params).fetchone()

        # Savings estimate: for class A, P75 exceeds median by avg %
        class_a = [d for d in position_details if d["conf_class"].startswith("A")]
        avg_savings_pct = 0
        if class_a:
            ratios = [(d["p75"] / d["median_price"] - 1) * 100 for d in class_a if d["median_price"] > 0]
            avg_savings_pct = round(sum(ratios) / len(ratios), 1) if ratios else 0

        # Total volume
        total_volume = sum(p["total_sum"] or 0 for p in positions)

        return {
            "filters": {"year_from": year_from, "year_to": year_to, "tru_type": tru_type, "exclude_rnu": exclude_rnu},
            "total_positions": total_positions,
            "class_distribution": class_counts,
            "total_suppliers": supplier_stats["total"] if supplier_stats else 0,
            "rnu_suppliers": supplier_stats["rnu_count"] if supplier_stats else 0,
            "qualified_suppliers": (supplier_stats["total"] - supplier_stats["rnu_count"]) if supplier_stats else 0,
            "avg_savings_pct": avg_savings_pct,
            "total_volume": total_volume,
            "positions": position_details,
        }
    finally:
        db.close()

@app.get("/api/dashboard")
def get_dashboard(
    enstru: str = Query("192021.530.000001"),
    region: Optional[str] = Query(None),
    year_from: int = Query(2024), year_to: int = Query(2026),
    exclude_rnu: bool = Query(True),
    status: str = Query("EXECUTED"),
):
    """Single-product dashboard data."""
    db = get_db()
    try:
        joins = ["JOIN skk_plan_enstru pe ON c.plan_id = pe.plan_id"]
        where = ["pe.enstru = ?", "c.year BETWEEN ? AND ?", "c.status = ?"]
        params = [enstru, year_from, year_to, status]
        if region:
            where.append("(CAST(c.delivery_kato_code / 10000000 AS INTEGER) = CAST(? AS INTEGER))")
            params.append(region)

        prod = db.execute(
            "SELECT pe.name, de.tru_type FROM skk_plan_enstru pe LEFT JOIN skk_dim_enstru de ON pe.enstru = de.code WHERE pe.enstru = ? LIMIT 1",
            [enstru]).fetchone()
        product_name = prod["name"] if prod else enstru
        tru_type = TRU_TYPE_MAP.get(prod["tru_type"], "неизвестно") if prod and prod["tru_type"] else "неизвестно"

        sql = f"""
            SELECT c.contract_card_id, c.provider_bin, c.plan_id, c.amt, c.qty,
                   c.delivery_kato_code, c.year, c.status,
                   ROUND(c.amt / NULLIF(c.qty, 0), 2) as unit_price
            FROM skk_contract c {' '.join(joins)}
            WHERE {' AND '.join(where)} AND c.qty > 0 AND c.amt > 0
        """
        contracts = db.execute(sql, params).fetchall()
        if not contracts:
            return JSONResponse({"error": "No contracts found", "count": 0}, status_code=404)

        prices = [r["unit_price"] for r in contracts if r["unit_price"] and r["unit_price"] > 0]
        if len(prices) < 3:
            return JSONResponse({"error": "Insufficient data", "count": len(prices)}, status_code=400)

        stats = _compute_stats(prices)
        corridor = 0.20
        corr_low = round(stats["median"] * (1 - corridor), 2)
        corr_high = round(stats["median"] * (1 + corridor), 2)

        price_reference = {
            "enstru": enstru, "name": product_name, "tru_type": tru_type,
            "n_items": stats["n"], "n_suppliers": len(set(r["provider_bin"] for r in contracts)),
            "n_contracts": len(contracts),
            "median_price": stats["median"], "mode_price": round(max(set(prices), key=prices.count), 2) if prices else None,
            "p25": stats["p25"], "p75": stats["p75"], "spread_pct": stats["spread"],
            "conf_class": stats["conf_class"], "corridor_low": corr_low, "corridor_high": corr_high,
            "national_share": 1.0, "regional_share": round(sum(1 for p in prices if corr_low <= p <= corr_high) / len(prices), 3),
            "price_basis": "региональный" if region else "национальный",
            "needs_segmentation": stats["spread"] > 150
        }

        # Regional breakdown
        region_rows = db.execute(f"""
            SELECT CAST(c.delivery_kato_code / 10000000 AS INTEGER) as region_code,
                   rk.full_name_ru as region_name, COUNT(*) as n_items,
                   ROUND(AVG(c.amt / NULLIF(c.qty, 0)), 2) as median_price
            FROM skk_contract c {' '.join(joins)}
            JOIN skk_ref_kato rk ON CAST(c.delivery_kato_code / 10000000 AS INTEGER) = CAST(rk.ab AS INTEGER)
            WHERE {' AND '.join(where)} AND c.qty > 0 AND c.amt > 0
            GROUP BY region_code ORDER BY n_items DESC
        """, params).fetchall()

        price_by_region = []
        for r in region_rows:
            rp = db.execute(f"""
                SELECT ROUND(c.amt / NULLIF(c.qty, 0), 2) as up FROM skk_contract c {' '.join(joins)}
                WHERE {' AND '.join(where)} AND c.qty > 0 AND c.amt > 0
                  AND CAST(c.delivery_kato_code / 10000000 AS INTEGER) = ? ORDER BY up
            """, params + [r["region_code"]]).fetchall()
            rprices = [x["up"] for x in rp if x["up"] and x["up"] > 0]
            if len(rprices) >= 3:
                rps = sorted(rprices)
                rp25, rp75 = rps[int(len(rps)*0.25)], rps[int(len(rps)*0.75)]
                in_corr = sum(1 for p in rprices if corr_low <= p <= corr_high) / len(rprices)
            else:
                rp25 = rp75 = r["median_price"]; in_corr = 1.0
            price_by_region.append({
                "enstru": enstru, "name": product_name,
                "region": r["region_name"] or f"Код {r['region_code']}",
                "n_items": r["n_items"], "median_price": r["median_price"],
                "p25": round(rp25,2), "p75": round(rp75,2),
                "corridor_low": corr_low, "corridor_high": corr_high,
                "in_corridor_share": round(in_corr, 3),
            })

        # Supplier registry
        supplier_rows = db.execute(f"""
            SELECT c.provider_bin, COALESCE(ds.name_ru, s.name_ru, c.provider_bin) as supplier_name,
                   COALESCE(ds.is_rnu, 0) as is_rnu, COALESCE(ds.kato, s.jur_kato_ru, '') as legal_kato,
                   COUNT(*) as n_contracts, COUNT(DISTINCT c.contract_card_id) as n_items,
                   SUM(CASE WHEN c.status='EXECUTED' THEN 1 ELSE 0 END) as n_exec_items,
                   SUM(c.amt) as total_sum
            FROM skk_contract c {' '.join(joins)}
            LEFT JOIN skk_dim_subject ds ON c.provider_bin = ds.bin_iin
            LEFT JOIN skk_subject s ON c.provider_bin = s.bin_iin
            WHERE {' AND '.join(where)}
            GROUP BY c.provider_bin ORDER BY total_sum DESC
        """, params).fetchall()

        supplier_registry = []
        for sr in supplier_rows:
            if exclude_rnu and sr["is_rnu"]: continue
            dr = db.execute(f"""
                SELECT DISTINCT rk.full_name_ru FROM skk_contract c {' '.join(joins)}
                JOIN skk_ref_kato rk ON CAST(c.delivery_kato_code / 10000000 AS INTEGER) = CAST(rk.ab AS INTEGER)
                WHERE {' AND '.join(where)} AND c.provider_bin = ? ORDER BY rk.full_name_ru LIMIT 10
            """, params + [sr["provider_bin"]]).fetchall()
            ya = db.execute(f"""
                SELECT MIN(c.year) as ymin, MAX(c.year) as ymax FROM skk_contract c {' '.join(joins)}
                WHERE {' AND '.join(where)} AND c.provider_bin = ?
            """, params + [sr["provider_bin"]]).fetchone()
            years_str = f"{ya['ymin']}" if ya["ymin"] == ya["ymax"] else f"{ya['ymin']}–{ya['ymax']}"
            supplier_registry.append({
                "supplier_bin": sr["provider_bin"], "supplier_name": sr["supplier_name"] or sr["provider_bin"],
                "is_rnu": bool(sr["is_rnu"]), "legal_region": sr["legal_kato"] or "",
                "delivery_regions": "; ".join(d["full_name_ru"] for d in dr if d["full_name_ru"]),
                "n_contracts": sr["n_contracts"], "n_items": sr["n_items"],
                "n_exec_items": sr["n_exec_items"], "total_sum": sr["total_sum"] or 0,
                "years_active": years_str,
            })

        year_dist = db.execute(f"""
            SELECT c.year, COUNT(*) as cnt FROM skk_contract c {' '.join(joins)}
            WHERE {' AND '.join(where)} GROUP BY c.year ORDER BY c.year
        """, params).fetchall()

        recent = db.execute(f"""
            SELECT c.contract_card_id as id, COALESCE(ds.name_ru, s.name_ru, c.provider_bin) as provider,
                   c.year, c.qty, c.amt, ROUND(c.amt / NULLIF(c.qty, 0), 2) as unit_price,
                   rk.full_name_ru as region, c.status
            FROM skk_contract c {' '.join(joins)}
            LEFT JOIN skk_dim_subject ds ON c.provider_bin = ds.bin_iin
            LEFT JOIN skk_subject s ON c.provider_bin = s.bin_iin
            LEFT JOIN skk_ref_kato rk ON CAST(c.delivery_kato_code / 10000000 AS INTEGER) = CAST(rk.ab AS INTEGER)
            WHERE {' AND '.join(where)}
            ORDER BY c.contract_card_id DESC LIMIT 30
        """, params).fetchall()

        return {
            "filters": {"enstru": enstru, "region": region, "year_from": year_from, "year_to": year_to},
            "product_name": product_name, "tru_type": tru_type,
            "total_volume": sum(r["amt"] for r in contracts),
            "total_quantity": sum(r["qty"] for r in contracts),
            "price_reference": [price_reference],
            "price_by_region": price_by_region,
            "supplier_registry": supplier_registry,
            "year_distribution": [{"year": r["year"], "count": r["cnt"]} for r in year_dist],
            "contracts": [{"id": str(r["id"]), "provider": r["provider"] or str(r["id"]),
                           "year": r["year"], "region": r["region"] or "—",
                           "qty": r["qty"], "amt": r["amt"], "unit_price": r["unit_price"],
                           "status": r["status"]} for r in recent],
        }
    finally:
        db.close()

@app.get("/api/report-docx")
def generate_report_docx(
    enstru: str = Query("192021.530.000001"),
    year_from: int = Query(2024), year_to: int = Query(2026),
):
    """Generate IBCS-compliant DOCX report for selected ENSTRU."""
    db = get_db()
    try:
        # Get dashboard data
        from urllib.request import urlopen
        import json

        # Call our own API internally
        host = "http://127.0.0.1:8100"
        try:
            resp = urlopen(f"{host}/api/dashboard?enstru={enstru}&year_from={year_from}&year_to={year_to}", timeout=30)
            data = json.loads(resp.read())
        except:
            return JSONResponse({"error": "Failed to fetch dashboard data"}, status_code=500)

        pr = data["price_reference"][0]
        regions = data["price_by_region"]
        suppliers = data["supplier_registry"]

        # Build DOCX
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # ── Title ──
        title = doc.add_heading('ОТЧЁТ', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading('о мониторинге рыночных цен ТРУ и формировании квалифицированного пула поставщиков', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Продукт: {pr['name']} ({pr['enstru']})")
        doc.add_paragraph(f"Период: {year_from}–{year_to} · Тип ТРУ: {pr['tru_type']} · Дата: {datetime.date.today().strftime('%d.%m.%Y')}")
        doc.add_paragraph()

        # ── 1. Резюме ──
        doc.add_heading('1. Резюме', level=1)
        doc.add_paragraph(
            f"Рекомендованная цена определена по {pr['n_items']} позициям с фактически исполненными договорами. "
            f"Класс достоверности: {pr['conf_class']}. "
            f"Медианная цена: {pr['median_price']} ₸/ед., диапазон P25–P75: {pr['p25']} – {pr['p75']} ₸/ед. "
            f"Ценовой спред: {pr['spread_pct']}%. "
            f"Квалифицированный пул: {pr['n_suppliers']} поставщиков с подтверждённым опытом."
        )

        # ── 2. Рекомендованные цены ──
        doc.add_heading('2. Рекомендованные цены', level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, text in enumerate(["Параметр", "Значение", "Параметр", "Значение", "", "", ""]):
            hdr[i].text = text if text else ""
        hdr[0].text = "ЕНС ТРУ"; hdr[1].text = pr['enstru']
        hdr[2].text = "Наименование"; hdr[3].text = pr['name'][:60]

        rows_data = [
            ("Количество позиций", str(pr['n_items']), "Поставщиков", str(pr['n_suppliers'])),
            ("Медианная цена", f"{pr['median_price']} ₸", "Мода", f"{pr['mode_price']} ₸"),
            ("P25", f"{pr['p25']} ₸", "P75", f"{pr['p75']} ₸"),
            ("Спред", f"{pr['spread_pct']}%", "Класс", pr['conf_class']),
            ("Коридор (±20%)", f"{pr['corridor_low']} – {pr['corridor_high']} ₸", "Базис", pr['price_basis']),
        ]
        for rd in rows_data:
            row = table.add_row().cells
            row[0].text = rd[0]; row[1].text = rd[1]
            row[2].text = rd[2]; row[3].text = rd[3]

        doc.add_paragraph()

        # ── 3. Региональный разрез ──
        doc.add_heading('3. Региональный разрез', level=1)
        if regions:
            reg_table = doc.add_table(rows=1, cols=6)
            reg_table.style = 'Light Grid Accent 1'
            reg_hdr = reg_table.rows[0].cells
            for i, t in enumerate(["Регион", "Позиций", "Медиана, ₸", "P25, ₸", "P75, ₸", "В коридоре"]):
                reg_hdr[i].text = t
            for r in regions:
                row = reg_table.add_row().cells
                row[0].text = r['region']; row[1].text = str(r['n_items'])
                row[2].text = str(r['median_price']); row[3].text = str(r['p25'])
                row[4].text = str(r['p75']); row[5].text = f"{r['in_corridor_share']*100:.0f}%"
        else:
            doc.add_paragraph("Недостаточно данных для регионального разреза.")
        doc.add_paragraph()

        # ── 4. Квалифицированный пул поставщиков ──
        doc.add_heading('4. Квалифицированный пул поставщиков', level=1)
        if suppliers:
            supp_table = doc.add_table(rows=1, cols=5)
            supp_table.style = 'Light Grid Accent 1'
            sh = supp_table.rows[0].cells
            for i, t in enumerate(["Поставщик", "Контрактов", "Сумма, ₸", "Годы", "Регионы поставок"]):
                sh[i].text = t
            for s in suppliers[:30]:
                row = supp_table.add_row().cells
                row[0].text = s['supplier_name'][:60]
                row[1].text = str(s['n_contracts'])
                row[2].text = f"{s['total_sum']:,.0f}"
                row[3].text = s['years_active']
                row[4].text = s['delivery_regions'][:80]
        else:
            doc.add_paragraph("Нет квалифицированных поставщиков.")
        doc.add_paragraph()

        # ── 5. Оговорки ──
        doc.add_heading('5. Оговорки по достоверности', level=1)
        caveats = [
            f"Класс достоверности: {pr['conf_class']}.",
            "Цены не приведены к дате оценки (инфляция/курс).",
            "Для позиций класса C требуется адресный запрос КП.",
            "Данные ограничены исполненными договорами СКК.",
        ]
        for c in caveats:
            doc.add_paragraph(c, style='List Bullet')

        # ── Save ──
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Отчёт_мониторинг_цен_{enstru}.docx"}
        )
    finally:
        db.close()

@app.get("/api/report")
def generate_report(enstru: str = Query("192021.530.000001")):
    """Trigger SKK pipeline report generation (Excel + DOCX)."""
    import subprocess, sys
    cmd = [sys.executable, "-c",
           f"from skk.intake import run_intake; run_intake(['{enstru}']); "
           "from skk.filter import run_filter; run_filter(); "
           "from skk.prices import run_prices; run_prices(); "
           "from skk.pool import run_pool; run_pool(); "
           "from skk.report import render; print(render())"]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(Path(__file__).parent), timeout=120)
    return {
        "success": result.returncode == 0,
        "output": result.stdout[-500:] if result.stdout else "",
        "error": result.stderr[-500:] if result.stderr else "",
    }

@app.get("/")
def serve_dashboard():
    dashboard_path = Path(__file__).parent / "skk-dashboard.html"
    if dashboard_path.exists():
        return FileResponse(dashboard_path, media_type="text/html")
    return {"message": "SKK Dashboard API v2"}
